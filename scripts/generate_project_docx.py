"""Generate PropAI.pk project documentation as a Word file (viva-ready)."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DESKTOP_OUTPUT = Path.home() / "Desktop" / "PropAI_Project_Documentation_Viva.docx"
DOWNLOADS_OUTPUT = Path.home() / "Downloads" / "PropAI_Project_Documentation_Viva.docx"
LEGACY_DESKTOP = Path.home() / "Desktop" / "PropAI_Project_Documentation.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text, level=0):
    return doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")


def add_code(doc, code: str, title: str | None = None):
    if title:
        add_para(doc, title, bold=True)
    p = doc.add_paragraph()
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # ── Title page ──────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("PropAI.pk\n")
    t_run.bold = True
    t_run.font.size = Pt(28)
    t_run.font.color.rgb = RGBColor(0x1A, 0x67, 0xF5)
    sub = title.add_run(
        "Predictive AI for Real Estate Investment Analysis\n"
        "Karachi & Islamabad, Pakistan\n\n"
        "Complete Project Documentation & Viva Guide\n"
        f"{date.today().strftime('%B %d, %Y')}"
    )
    sub.font.size = Pt(14)

    doc.add_page_break()

    # ── 1. Executive Summary ───────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "PropAI.pk is a full-stack web application for real estate investment analysis in Pakistan, "
        "focused on Karachi and Islamabad. It ingests Zameen.com listing data, stores it in PostgreSQL "
        "(Neon cloud or Docker local), trains ML models to predict property prices, and provides "
        "role-based dashboards for Investors, Agents, and Administrators."
    )
    add_para(doc, "Key capabilities:", bold=True)
    for item in [
        "Dynamic landing page with live platform statistics from GET /stats/platform",
        "Historical price statistics by area (SQL aggregates — no ML)",
        "ML-based price predictions with location risk scoring",
        "Agent listing form with city-filtered dropdowns (Marla/Kanal, area size, custom values)",
        "Property search, favorites, and listing management",
        "Market analytics with Recharts visualizations",
        "JWT authentication with Role-Based Access Control (RBAC)",
        "Docker Compose deployment for one-command setup",
        "Admin CSV dataset upload with validation",
    ]:
        add_bullet(doc, item)

    # ── 2. Technology Stack ────────────────────────────────────────────────
    add_heading(doc, "2. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technologies"],
        [
            ["Backend", "FastAPI, SQLAlchemy, Alembic, Pydantic, JWT, bcrypt"],
            ["Database", "Neon PostgreSQL (SSL) or Docker Postgres 16"],
            ["Machine Learning", "scikit-learn, XGBoost, pandas, joblib"],
            ["Frontend", "React 19, Vite, TypeScript, Tailwind CSS v4"],
            ["State & UI", "Zustand, Recharts, Framer Motion, react-hot-toast"],
            ["Deployment", "Docker Compose (db + backend + frontend)"],
            ["Data Source", "zameen-updated.csv (~168K rows; ~97K Karachi + Islamabad)"],
        ],
    )

    # ── 3. Project Structure ───────────────────────────────────────────────
    add_heading(doc, "3. Project Structure", 1)
    add_para(
        doc,
        """ai project/
├── backend/
│   ├── app/
│   │   ├── main.py              # FastAPI entry, CORS, routers
│   │   ├── config.py            # SECRET_KEY, DATABASE_URL (backend only)
│   │   ├── dependencies.py      # JWT auth + require_roles() RBAC
│   │   ├── routers/             # auth, users, properties, predictions, analytics, stats
│   │   ├── services/            # Business logic
│   │   ├── models/              # User, Property, Favorite
│   │   └── ml/                  # preprocessing, train, inference, artifacts/
│   ├── alembic/                 # DB migrations
│   └── scripts/                 # seed_db, inspect_dataset, test_connection
├── frontend/
│   └── src/
│       ├── pages/               # Landing, Auth, Dashboards, AgentListingForm
│       ├── components/          # investor/, layout/, cyber/, ui/
│       ├── hooks/               # usePlatformStats (live landing data)
│       ├── services/            # API wrappers
│       └── store/               # Zustand auth (JWT token only)
├── data/                        # Writable dataset folder (Docker uploads)
├── docker-compose.yml
└── zameen-updated.csv"""
    )

    # ── 4. Database Schema ─────────────────────────────────────────────────
    add_heading(doc, "4. Database Schema", 1)
    add_para(doc, "UI label 'Area' maps to database column location (e.g. DHA Defence, G-10).")
    add_table(
        doc,
        ["Table", "Key Columns", "Purpose"],
        [
            ["users", "email, hashed_password, role, is_active", "Authentication & RBAC"],
            ["properties", "location, city, price, property_type, area_type, area_size, agent_id", "Listings"],
            ["favorites", "user_id, property_id", "Investor saved properties"],
        ],
    )
    add_table(
        doc,
        ["UI Label", "DB Column", "Example"],
        [
            ["Area", "location", "DHA Phase 6"],
            ["City", "city", "Karachi"],
            ["Size unit", "area_type", "Marla, Kanal"],
            ["Size value", "area_size", "5, 10, 1.5"],
            ["Category", "area_category", "5-10 Marla"],
        ],
    )

    # ── 5. Dataset ─────────────────────────────────────────────────────────
    add_heading(doc, "5. Dataset Details", 1)
    add_table(
        doc,
        ["Attribute", "Value"],
        [
            ["File", "zameen-updated.csv (also copy in data/ for Docker)"],
            ["Total rows (raw)", "~168,446"],
            ["Filtered for ML", "Karachi + Islamabad only (~97,910 rows)"],
            ["Price format", "Numeric PKR; pipeline also parses Lakh/Crore text"],
            ["Cities used", "Karachi, Islamabad"],
        ],
    )
    add_para(
        doc,
        "IMPORTANT FOR VIVA: If only the 8-row sample CSV was uploaded via Admin Dataset page, "
        "the model will be trained on 8 rows and Random Forest accuracy will drop sharply. "
        "Restore the full zameen-updated.csv and re-run train + seed for production-quality metrics.",
        bold=True,
    )

    # ── 6. ML Pipeline ─────────────────────────────────────────────────────
    add_heading(doc, "6. Machine Learning Pipeline", 1)
    add_para(doc, "Training: cd backend && python -m app.ml.train", bold=True)
    add_para(
        doc,
        "One unified model is trained on ALL property types (House, Flat, Plot) together — "
        "NOT separate models per type or per area. property_type and location are input features."
    )

    add_heading(doc, "6.1 Expected Metrics (Full ~97K Dataset)", 2)
    add_table(
        doc,
        ["Model", "R²", "MAE (PKR)", "RMSE (PKR)", "Notes"],
        [
            ["Linear Regression", "~45%", "~14.2M", "~27.3M", "Baseline"],
            ["Random Forest", "~90%", "~3.4M", "~11.7M", "Best on full data — selected"],
            ["XGBoost", "~90%", "~3.8M", "~11.8M", "Close second"],
        ],
    )

    add_heading(doc, "6.2 ML Features Used", 2)
    add_table(
        doc,
        ["Type", "Features"],
        [
            ["Categorical (One-Hot)", "property_type, location, city, area_type, area_category, purpose"],
            ["Numeric", "baths, bedrooms, area_size"],
            ["Target", "price_numeric (PKR)"],
        ],
    )

    add_heading(doc, "6.3 Risk Score (Not ML)", 2)
    add_para(
        doc,
        "Risk score measures price volatility per city+location: volatility = std(price) / mean(price). "
        "Mapped to 0–100 and labeled low / medium / high. It is NOT investment quality — it shows "
        "how widely prices vary in historical listings for that area."
    )

    add_heading(doc, "6.4 Artifacts", 2)
    add_bullet(doc, "backend/app/ml/artifacts/model.joblib — best pipeline (preprocessor + model)")
    add_bullet(doc, "metrics.json — R², MAE, RMSE per model")
    add_bullet(doc, "location_volatility.csv — risk score lookup table")

    # ── 7. Frontend Architecture ─────────────────────────────────────────
    add_heading(doc, "7. Frontend Architecture", 1)
    add_table(
        doc,
        ["Page", "Route", "Role", "Description"],
        [
            ["Landing", "/", "Public", "Cyber-themed; live stats from /stats/platform"],
            ["Login / Register", "/login, /register", "Guest", "Split-panel auth UI"],
            ["Investor Dashboard", "/investor", "Investor", "Browse by Area + Predict Investment"],
            ["Market Analytics", "/investor/analytics", "Investor", "Recharts city/area charts"],
            ["Favorites", "/investor/favorites", "Investor", "Saved properties"],
            ["Agent Dashboard", "/agent", "Agent", "Listing table + stats"],
            ["Add/Edit Listing", "/agent/new, /agent/edit/:id", "Agent", "Cascading dropdowns"],
            ["Admin Overview", "/admin", "Admin", "System stats + ML metrics table"],
            ["User Management", "/admin/users", "Admin", "Activate/deactivate users"],
            ["Dataset Upload", "/admin/dataset", "Admin", "CSV validation + save"],
        ],
    )

    add_heading(doc, "7.1 Three Separate Data Paths (Critical for Viva)", 2)
    add_table(
        doc,
        ["Section", "Data Source", "API", "Uses ML?"],
        [
            ["Browse by Area", "Postgres SQL aggregates", "/stats/areas, /stats/areas/prices", "No"],
            ["Predict Investment", "joblib model", "POST /predictions", "Yes"],
            ["Market Analytics", "Postgres grouped stats", "/stats/cities, /stats/cities/{city}/areas", "No"],
        ],
    )

    add_heading(doc, "7.2 Agent Listing Form Features", 2)
    for item in [
        "Property type dropdown from GET /stats/property-types",
        "City selection filters area/location dropdown via GET /stats/areas?city=",
        "Marla/Kanal selection filters area category and size options",
        "ComboSelect component: pick from list OR 'Other' for custom manual entry",
        "Validation before POST /properties",
    ]:
        add_bullet(doc, item)

    # ── 8. Authentication & Security ───────────────────────────────────────
    add_heading(doc, "8. Authentication & Security", 1)
    add_para(
        doc,
        "JWT Bearer tokens issued by backend. Passwords hashed with bcrypt. "
        "SECRET_KEY and DATABASE_URL exist ONLY in backend/.env — never in frontend code."
    )
    add_table(
        doc,
        ["Item", "Where Stored", "Safe?"],
        [
            ["SECRET_KEY", "backend/.env only", "Yes — server-side"],
            ["DATABASE_URL", "backend/.env only", "Yes — server-side"],
            ["JWT access token", "Browser localStorage (Zustand)", "User session — not a server secret"],
            ["VITE_API_URL", "frontend/.env", "Yes — public API URL only"],
        ],
    )
    add_table(
        doc,
        ["Demo Account", "Email", "Password", "Role"],
        [
            ["Admin", "admin@realestate.pk", "admin123", "admin"],
            ["Investor", "investor@realestate.pk", "investor123", "investor"],
            ["Agent", "agent@realestate.pk", "agent123", "agent"],
        ],
    )

    # ── 9. API Endpoints ───────────────────────────────────────────────────
    add_heading(doc, "9. API Endpoints Reference", 1)
    add_table(
        doc,
        ["Prefix", "Key Endpoints", "Roles"],
        [
            ["/auth", "POST /register, /login/json, GET /me", "All"],
            ["/stats", "GET /platform (public), /areas, /areas/prices, /property-types", "Varies"],
            ["/predictions", "POST /predictions, GET /property/{id}", "Authenticated"],
            ["/properties", "GET search, POST, PATCH, DELETE, favorites", "Varies"],
            ["/analytics", "GET /system, /model, /agent, POST /dataset/upload", "Admin/Agent/Investor"],
            ["/users", "GET /users, PATCH /users/{id}", "Admin"],
        ],
    )

    # ── 10. Docker Deployment ──────────────────────────────────────────────
    add_heading(doc, "10. Docker Deployment", 1)
    add_para(doc, "Command: docker compose up --build", bold=True)
    for item in [
        "db — Postgres 16 on port 5432",
        "backend — auto migrations, ML train if needed, seed on start",
        "frontend — nginx on port 5173",
        "Dataset uploads write to ./data folder (writable volume)",
        "After admin CSV upload: docker compose exec backend python -m app.ml.train",
    ]:
        add_bullet(doc, item)

    # ── 11. Setup Workflow ─────────────────────────────────────────────────
    add_heading(doc, "11. Setup & Installation (Local)", 1)
    for step in [
        "Copy .env.example → backend/.env (DATABASE_URL + SECRET_KEY)",
        "cd backend && pip install -r requirements.txt",
        "python -m alembic upgrade head",
        "python -m app.ml.train",
        "python scripts/seed_db.py",
        "cd frontend && npm install && npm run dev",
        "Backend: python -m uvicorn app.main:app --reload --port 8000",
    ]:
        add_bullet(doc, step)

    # ── 12. RBAC Matrix ────────────────────────────────────────────────────
    add_heading(doc, "12. Role-Based Access Control Matrix", 1)
    add_table(
        doc,
        ["Capability", "Admin", "Investor", "Agent"],
        [
            ["ML price prediction", "Yes", "Yes", "No"],
            ["Browse by Area (stats)", "Yes", "Yes", "No"],
            ["Market Analytics", "Yes", "Yes", "No"],
            ["Favorites", "No", "Yes", "No"],
            ["Add/edit own listings", "Yes", "No", "Yes"],
            ["User management", "Yes", "No", "No"],
            ["Upload CSV dataset", "Yes", "No", "No"],
            ["View ML metrics", "Yes", "Yes", "No"],
        ],
    )

    # ── 13. CODE SNIPPETS FOR VIVA ─────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "13. Key Code Snippets (Viva Reference)", 1)
    add_para(
        doc,
        "Instructors often ask you to explain specific code. Below are the most important snippets "
        "with file paths you should know."
    )

    add_heading(doc, "13.1 ML Preprocessor & Model Training", 2)
    add_code(
        doc,
        '''# backend/app/ml/train.py
def _build_preprocessor() -> ColumnTransformer:
    categorical = ["property_type", "location", "city",
                   "area_type", "area_category", "purpose"]
    numeric = ["baths", "bedrooms", "area_size"]
    return ColumnTransformer([
        ("cat", OneHotEncoder(handle_unknown="ignore"), categorical),
        ("num", "passthrough", numeric),
    ])

# Train 3 models, pick highest R² on 20% test split
pipeline = Pipeline([("prep", preprocessor), ("model", estimator)])
pipeline.fit(X_train, y_train)
joblib.dump(best_pipeline, "model.joblib")''',
        "Feature engineering + model selection:",
    )

    add_heading(doc, "13.2 Price Prediction & Risk Score", 2)
    add_code(
        doc,
        '''# backend/app/ml/inference.py
def predict_price(features: dict) -> float:
    model = joblib.load("model.joblib")
    return float(max(0, model.predict(pd.DataFrame([features]))[0]))

def compute_risk_score(city, location):
    volatility = std(price) / mean(price)  # per area from CSV
    risk_score = min(100, volatility * 100)
    # <33 = low, <66 = medium, else high''',
        "Inference at runtime:",
    )

    add_heading(doc, "13.3 JWT Authentication & RBAC", 2)
    add_code(
        doc,
        '''# backend/app/dependencies.py
def require_roles(*roles: UserRole):
    def role_checker(current_user: User = Depends(get_current_user)):
        if current_user.role not in roles:
            raise HTTPException(403, "Not authorized")
        return current_user
    return role_checker

# Usage in router:
@router.get("/system")
def system_analytics(_: User = Depends(require_roles(UserRole.ADMIN))):
    ...''',
        "Role guard pattern:",
    )

    add_heading(doc, "13.4 Password Hashing (bcrypt)", 2)
    add_code(
        doc,
        '''# backend/app/services/auth_service.py
def get_password_hash(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def create_access_token(user_id, role):
    payload = {"sub": str(user_id), "role": role, "exp": ...}
    return jwt.encode(payload, settings.SECRET_KEY, algorithm="HS256")''',
        "Secrets stay in backend config only:",
    )

    add_heading(doc, "13.5 Public Platform Stats (Dynamic Landing Page)", 2)
    add_code(
        doc,
        '''# backend/app/routers/stats.py — NO auth required
@router.get("/platform")
def platform_stats(db: Session = Depends(get_db)):
    return svc.get_platform_stats(db)  # live DB counts + metrics.json

# frontend/src/types/platform.ts
const totalListings = cities.reduce((sum, c) => sum + c.property_count, 0);
const modelAccuracyPct = model.r2 != null ? Math.round(model.r2 * 1000) / 10 : null;''',
        "Landing page numbers are computed from API — not hardcoded:",
    )

    add_heading(doc, "13.6 Frontend Route Protection", 2)
    add_code(
        doc,
        '''// frontend/src/components/RouteGuard.tsx
if (!token) return <Navigate to="/login" />;
if (allowedRoles && !allowedRoles.includes(user.role))
    return <Navigate to={getDashboardPath(user.role)} />;
return children;''',
        "Protects /investor, /agent, /admin routes:",
    )

    add_heading(doc, "13.7 Dataset Validation on Upload", 2)
    add_code(
        doc,
        '''# backend/app/routers/analytics.py
df = pd.read_csv(uploaded_file)
report = validate_dataset(df)  # checks required columns, Karachi/Islamabad
dest = Path(settings.DATASET_PATH)
dest.write_bytes(contents)  # saves to data/zameen-updated.csv''',
        "Admin dataset upload flow:",
    )

    add_heading(doc, "13.8 Agent Form — City Filters Areas", 2)
    add_code(
        doc,
        '''// frontend/src/pages/AgentListingForm.tsx
useEffect(() => {
  statsService.getAreas(form.city).then(areas => {
    setAreaOptions([...new Set(areas.map(a => a.location))]);
  });
}, [form.city]);

// ComboSelect: dropdown + "Other (enter manually)" text input''',
        "Cascading dropdowns:",
    )

    # ── 14. VIVA QUESTIONS & ANSWERS ──────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "14. Viva Questions & Model Answers", 1)

    qa = [
        (
            "What is PropAI.pk and what problem does it solve?",
            "It helps investors and agents analyze Karachi/Islamabad property markets using "
            "historical listing statistics and ML price predictions, reducing guesswork in investment decisions.",
        ),
        (
            "Why Random Forest and not a single linear model?",
            "Property prices have non-linear relationships (location, type, size). Random Forest "
            "achieved ~90% R² vs ~45% for Linear Regression on the full dataset. We compare 3 models "
            "and auto-select the best by R².",
        ),
        (
            "Is the model trained separately for flats vs houses?",
            "No. One unified model uses property_type as a categorical feature alongside location, "
            "city, size, beds, and baths. Future improvement: separate models per type.",
        ),
        (
            "What is the difference between Browse by Area and Predict Investment?",
            "Browse by Area uses SQL aggregates from Postgres (min/max/avg) — no ML. "
            "Predict Investment calls the trained model via POST /predictions. They are intentionally separate.",
        ),
        (
            "How is risk score calculated?",
            "Risk = price volatility (std/mean) per city+location from training data, scaled 0–100. "
            "High risk means wide price spread in that area, not necessarily a bad investment.",
        ),
        (
            "Explain your authentication flow.",
            "User logs in → backend validates bcrypt hash → issues JWT → frontend stores token in "
            "Zustand/localStorage → sends Bearer token on each API call → require_roles() checks role.",
        ),
        (
            "Where are secrets stored?",
            "SECRET_KEY and DATABASE_URL only in backend/.env. Frontend has VITE_API_URL (public URL only). "
            "No API keys or DB passwords in React code.",
        ),
        (
            "What is R², MAE, and RMSE?",
            "R² = how much variance the model explains (1.0 = perfect). MAE = average error in PKR. "
            "RMSE = penalizes large errors more. On full data MAE ~3.4M PKR (~34 Lakh average error).",
        ),
        (
            "Why might Random Forest show low accuracy in Admin dashboard?",
            "If the 8-row sample CSV replaced the full dataset, the model trains on only 8 properties. "
            "Restore full zameen-updated.csv and run: python -m app.ml.train",
        ),
        (
            "What does the Admin Dataset upload do?",
            "Validates CSV schema, saves to data/zameen-updated.csv. Admin must then retrain ML and "
            "optionally re-seed Postgres.",
        ),
        (
            "How does Docker deployment work?",
            "docker-compose.yml runs Postgres + FastAPI backend + React frontend. Entrypoint runs "
            "migrations, trains model if missing, seeds DB, starts uvicorn.",
        ),
        (
            "What is OneHotEncoder handle_unknown='ignore'?",
            "If prediction input has a location not seen in training, encoder ignores it instead of "
            "crashing — important for new areas.",
        ),
        (
            "Explain RBAC in your project.",
            "Three roles: admin (users, dataset, system stats), investor (browse, predict, favorites), "
            "agent (own listings only). Enforced in backend via require_roles() and frontend RouteGuard.",
        ),
        (
            "What is the landing page live data?",
            "GET /stats/platform returns total listings, cities, property types, best model R² from "
            "live DB + metrics.json. Frontend derives hero numbers with .reduce() — updates when data changes.",
        ),
    ]

    for i, (q, a) in enumerate(qa, 1):
        add_para(doc, f"Q{i}: {q}", bold=True)
        add_para(doc, f"A: {a}")
        doc.add_paragraph()

    # ── 15. Demo Script for Viva ────────────────────────────────────────────
    add_heading(doc, "15. Live Demo Script (5–10 minutes)", 1)
    steps = [
        ("1", "Open http://localhost:5173 — show landing page live stats (listings, R², cities)"),
        ("2", "Register as Investor or login investor@realestate.pk / investor123"),
        ("3", "Investor Dashboard → Browse by Area: pick DHA/Karachi, show price slider from real min/max"),
        ("4", "Predict Investment: enter property details → show ML price + risk score"),
        ("5", "Market Analytics → city chart + area drill-down"),
        ("6", "Logout → login agent@realestate.pk / agent123"),
        ("7", "Agent → Add Listing: show city→area cascade, Marla/Kanal dropdowns"),
        ("8", "Login admin@realestate.pk / admin123 → Admin Overview ML metrics table"),
        ("9", "Show API docs at http://localhost:8000/docs"),
    ]
    for num, desc in steps:
        add_bullet(doc, f"Step {num}: {desc}")

    # ── 16. Limitations & Future Work ──────────────────────────────────────
    add_heading(doc, "16. Limitations & Future Work", 1)
    for item in [
        "Single model for all property types — separate models would improve accuracy",
        "Static model — no automatic retraining when market changes",
        "Risk score is volatility only — not legal, liquidity, or demand analysis",
        "MAE ~34 Lakh on full data — estimates not exact valuations",
        "Future: confidence intervals, per-city models, httpOnly cookies, model monitoring",
    ]:
        add_bullet(doc, item)

    # ── Footer ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_para(doc, f"Generated: {date.today().strftime('%B %d, %Y')}")
    add_para(doc, "Project: PropAI.pk — FYP / Final Year Project")
    add_para(doc, "Authors: [Add your name and student ID here]")
    add_para(doc, "Supervisor: [Add supervisor name here]")
    add_para(doc, f"Source code: {PROJECT_ROOT}")

    outputs = [DESKTOP_OUTPUT, DOWNLOADS_OUTPUT]
    if not LEGACY_DESKTOP.exists() or LEGACY_DESKTOP.stat().st_size == 0:
        outputs.append(LEGACY_DESKTOP)
    for path in outputs:
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(path)
            print(f"Saved: {path}")
        except PermissionError:
            print(f"Skipped (file open?): {path}")


if __name__ == "__main__":
    build()
