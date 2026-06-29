"""Generate PropAI.pk project documentation as a Word file."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path.home() / "Downloads" / "PropAI_Project_Documentation.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text, level=0):
    return doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("PropAI.pk\n")
    t_run.bold = True
    t_run.font.size = Pt(28)
    t_run.font.color.rgb = RGBColor(0x1A, 0x67, 0xF5)
    sub = title.add_run(
        "Predictive AI for Real Estate Investment Analysis\n"
        "Karachi & Islamabad, Pakistan\n\n"
        f"Complete Project Documentation\n{date.today().strftime('%B %d, %Y')}"
    )
    sub.font.size = Pt(14)

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "PropAI.pk is a full-stack web application designed for real estate investment analysis "
        "in Pakistan, focusing on Karachi and Islamabad. The platform ingests property listing data "
        "(Zameen.com CSV), stores it in a Neon PostgreSQL database, trains machine learning models "
        "to predict property prices, and exposes role-based dashboards for Investors, Agents, and Administrators."
    )
    add_para(doc, "Key capabilities:", bold=True)
    for item in [
        "Historical price statistics by area (SQL aggregates — no ML)",
        "ML-based price predictions with location risk scoring",
        "Property search, favorites, and listing management",
        "Market analytics with interactive Recharts visualizations",
        "JWT authentication with Role-Based Access Control (RBAC)",
    ]:
        add_bullet(doc, item)

    # 2. Technology Stack
    add_heading(doc, "2. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technologies"],
        [
            ["Backend", "FastAPI, SQLAlchemy, Alembic, Pydantic, JWT, bcrypt"],
            ["Database", "Neon PostgreSQL (SSL, pooled connections)"],
            ["Machine Learning", "scikit-learn, XGBoost, pandas, joblib"],
            ["Frontend", "React 19, Vite, TypeScript, Tailwind CSS v4"],
            ["State & UI", "Zustand, Recharts, Framer Motion, react-hot-toast"],
            ["Data Source", "zameen-updated.csv (168K+ rows, filtered to ~97K for KHI/ISB)"],
        ],
    )

    # 3. Project Structure
    add_heading(doc, "3. Project Structure", 1)
    structure = """
ai project/
├── backend/                    # FastAPI application
│   ├── app/
│   │   ├── main.py             # App entry point, CORS, routers
│   │   ├── config.py           # Environment settings
│   │   ├── database.py         # SQLAlchemy + Neon connection
│   │   ├── dependencies.py     # Auth guards, RBAC
│   │   ├── models/             # User, Property, Favorite ORM models
│   │   ├── schemas/            # Pydantic request/response schemas
│   │   ├── routers/            # auth, users, properties, predictions, analytics, stats
│   │   ├── services/           # Business logic layer
│   │   └── ml/                 # Preprocessing, training, inference
│   ├── alembic/                # Database migrations
│   ├── scripts/                # test_connection, seed_db, inspect_dataset
│   └── requirements.txt
├── frontend/                   # React + Vite application
│   └── src/
│       ├── pages/              # Landing, Auth, Dashboards, Analytics
│       ├── components/         # UI, charts, investor sections
│       ├── services/           # API client wrappers
│       └── store/              # Zustand auth state
├── zameen-updated.csv          # Source dataset
├── .env.example                # Environment variable template
└── README.md
"""
    add_para(doc, structure)

    # 4. Database Schema
    add_heading(doc, "4. Database Schema", 1)
    add_para(
        doc,
        "The PostgreSQL database contains three main tables. In the UI, 'Area' maps to the "
        "location column (neighborhood/sector names such as DHA Defence, G-10)."
    )
    add_table(
        doc,
        ["Table", "Key Columns", "Purpose"],
        [
            ["users", "email, full_name, hashed_password, role, is_active", "Authentication & RBAC"],
            ["properties", "location, city, price, property_type, bedrooms, baths, area_type, area_size, agent_id", "Listings from CSV + agent-created"],
            ["favorites", "user_id, property_id", "Investor saved properties"],
        ],
    )
    add_para(doc, "User roles (enum): admin, investor, agent", bold=True)

    # 5. Dataset
    add_heading(doc, "5. Dataset Details", 1)
    add_table(
        doc,
        ["Attribute", "Value"],
        [
            ["File", "zameen-updated.csv"],
            ["Total rows", "168,446"],
            ["Cities in raw data", "Karachi, Islamabad, Lahore, Rawalpindi, Faisalabad"],
            ["Filtered for ML/seed", "Karachi (60,484) + Islamabad (37,426) = 97,910"],
            ["Price format", "Numeric PKR (integers); pipeline also supports Lakh/Crore text"],
            ["Size units", "Marla, Kanal (area_type + area_size columns)"],
        ],
    )
    add_para(doc, "Key CSV columns:", bold=True)
    add_bullet(doc, "property_id, property_type, price, location, city, bedrooms, baths")
    add_bullet(doc, "Area Type → area_type, Area Size → area_size, Area Category → area_category")

    # 6. ML Pipeline
    add_heading(doc, "6. Machine Learning Pipeline", 1)
    add_para(doc, "Training command: python -m app.ml.train (run from /backend)", bold=True)
    add_para(
        doc,
        "The pipeline cleans data, trains three models, and saves the best performer:"
    )
    add_table(
        doc,
        ["Model", "R² Score", "MAE (PKR)", "RMSE (PKR)"],
        [
            ["Linear Regression (baseline)", "44.9%", "14.2 Million", "27.3 Million"],
            ["Random Forest (selected)", "90.0%", "3.4 Million", "11.7 Million"],
            ["XGBoost", "89.7%", "3.8 Million", "11.8 Million"],
        ],
    )
    add_para(doc, "Artifacts saved to backend/ml/artifacts/:", bold=True)
    add_bullet(doc, "model.joblib — trained Random Forest pipeline")
    add_bullet(doc, "metrics.json — evaluation metrics for Admin dashboard")
    add_bullet(doc, "location_volatility.csv — used for risk score calculation")
    add_para(
        doc,
        "Risk score: based on price volatility (standard deviation / mean) per city+location. "
        "Scores 0–100 mapped to low / medium / high risk levels."
    )

    # 7. Setup Workflow
    add_heading(doc, "7. Setup & Installation Workflow", 1)
    steps = [
        ("Step 1", "Copy .env.example to backend/.env and add Neon DATABASE_URL (pooled) and SECRET_KEY"),
        ("Step 2", "cd backend && pip install -r requirements.txt"),
        ("Step 3", "python scripts/test_connection.py — verify Neon connectivity"),
        ("Step 4", "python -m alembic upgrade head — create database tables"),
        ("Step 5", "python -m app.ml.train — train ML model on CSV"),
        ("Step 6", "python scripts/seed_db.py — load properties and demo users"),
        ("Step 7", "cd frontend && npm install"),
    ]
    for step, desc in steps:
        add_bullet(doc, f"{step}: {desc}")

    add_heading(doc, "7.1 Running the Application", 2)
    add_para(doc, "Terminal 1 (Backend):", bold=True)
    add_para(doc, "cd backend\npython -m uvicorn app.main:app --reload --port 8000")
    add_para(doc, "Terminal 2 (Frontend):", bold=True)
    add_para(doc, "cd frontend\nnpm run dev")
    add_para(doc, "URLs:", bold=True)
    add_bullet(doc, "Frontend: http://localhost:5173")
    add_bullet(doc, "API Docs: http://localhost:8000/docs")
    add_bullet(doc, "Health: http://localhost:8000/health")

    # 8. Authentication
    add_heading(doc, "8. Authentication & Security", 1)
    add_para(
        doc,
        "JWT-based authentication with bcrypt password hashing. Tokens are stored in Zustand "
        "with localStorage persistence. RouteGuard component protects frontend routes by role."
    )
    add_table(
        doc,
        ["Demo Account", "Email", "Password", "Role"],
        [
            ["Admin", "admin@realestate.pk", "admin123", "admin"],
            ["Investor", "investor@realestate.pk", "investor123", "investor"],
            ["Agent", "agent@realestate.pk", "agent123", "agent"],
        ],
    )

    # 9. API Endpoints
    add_heading(doc, "9. API Endpoints Reference", 1)
    add_table(
        doc,
        ["Prefix", "Key Endpoints", "Roles"],
        [
            ["/auth", "POST /register, /login/json, GET /me", "All"],
            ["/users", "GET /users, PATCH /users/{id}, DELETE /users/{id}", "Admin"],
            ["/properties", "GET search, POST create, PATCH/DELETE, favorites", "All (varies)"],
            ["/stats", "GET /areas, /areas/prices, /cities, /cities/{city}/areas", "Investor, Admin"],
            ["/predictions", "POST /predictions, GET /property/{id}", "Investor, Admin"],
            ["/analytics", "GET /system, /model, /agent, POST /dataset/upload", "Admin, Agent, Investor"],
        ],
    )

    # 10. Role-Based Workflows
    add_heading(doc, "10. Role-Based Workflows", 1)

    add_heading(doc, "10.1 Investor Role", 2)
    add_para(doc, "Purpose: Research properties, analyze areas, get ML estimates, save favorites.")
    add_table(
        doc,
        ["Page", "Route", "Description"],
        [
            ["Dashboard", "/investor", "Browse by Area + Predict My Investment"],
            ["Market Analytics", "/investor/analytics", "City comparison & area drill-down charts"],
            ["Favorites", "/investor/favorites", "Saved property listings"],
            ["Property Detail", "/properties/:id", "Listing info + predicted vs listed price + risk"],
        ],
    )
    add_para(doc, "Section A — Browse by Area (Statistics, NO ML):", bold=True)
    for step in [
        "Select area from dropdown (populated from database via GET /stats/areas)",
        "System fetches min, max, average price for that area (GET /stats/areas/prices)",
        "Dual-handle price slider appears, bounded to real min/max for that area",
        "Bedroom breakdown chips show average price per bedroom count",
        "Property list updates as slider moves (GET /properties with exact location + price range)",
        "User can save properties to favorites",
    ]:
        add_bullet(doc, step)

    add_para(doc, "Section B — Predict My Investment (ML ONLY):", bold=True)
    for step in [
        "Separate green-accent card — visually distinct from statistics section",
        "User selects: Area, Property Type, Bedrooms, Bathrooms, Size + unit (Marla/Kanal)",
        "On submit: POST /predictions calls trained Random Forest model",
        "Displays estimated price, risk score, and disclaimer text",
        "Does NOT share state with Browse by Area section",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "10.2 Agent Role", 2)
    add_para(doc, "Purpose: Manage own property listings and view performance analytics.")
    add_table(
        doc,
        ["Page", "Route", "Description"],
        [
            ["Agent Dashboard", "/agent", "Table of own listings + summary stats"],
            ["Add Listing", "/agent/new", "Create new property in Karachi or Islamabad"],
            ["Edit Listing", "/agent/edit/:id", "Update own listing only"],
        ],
    )
    for rule in [
        "Can only create listings in Karachi or Islamabad",
        "Can only edit/delete listings where agent_id matches logged-in user",
        "Agent analytics: total listings, active listings, average listed price (GET /analytics/agent)",
        "New listings stored in Postgres separately from CSV seed data",
    ]:
        add_bullet(doc, rule)

    add_heading(doc, "10.3 Admin Role", 2)
    add_para(doc, "Purpose: System management, user control, dataset uploads, ML metrics.")
    add_table(
        doc,
        ["Page", "Route", "Description"],
        [
            ["Admin Overview", "/admin", "System stats + ML model performance table"],
            ["User Management", "/admin/users", "View, activate, deactivate users"],
            ["Dataset Management", "/admin/dataset", "Upload and validate new CSV files"],
        ],
    )
    add_para(doc, "Dataset refresh workflow after CSV upload:", bold=True)
    for step in [
        "Upload CSV via Admin Dataset page (validates schema before saving)",
        "Run: python -m app.ml.train (retrain model)",
        "Run: python scripts/seed_db.py (reload DB if needed — skips if data exists)",
        "Check Admin Overview for updated model metrics",
    ]:
        add_bullet(doc, step)

    # 11. RBAC Matrix
    add_heading(doc, "11. Role-Based Access Control Matrix", 1)
    add_table(
        doc,
        ["Capability", "Admin", "Investor", "Agent"],
        [
            ["Search/browse properties", "Yes", "Yes", "Own listings"],
            ["ML price prediction", "Yes", "Yes", "No"],
            ["Browse by Area (stats)", "Yes", "Yes", "No"],
            ["Market Analytics charts", "Yes", "Yes", "No"],
            ["Save favorites", "No", "Yes", "No"],
            ["Add/edit listings", "Yes", "No", "Own only"],
            ["User management", "Yes", "No", "No"],
            ["Upload CSV dataset", "Yes", "No", "No"],
            ["System-wide analytics", "Yes", "No", "No"],
            ["Agent listing analytics", "No", "No", "Yes"],
        ],
    )

    # 12. Data Flow
    add_heading(doc, "12. Data Flow Architecture", 1)
    add_para(doc, "Three separate data paths (intentionally not merged):", bold=True)
    add_table(
        doc,
        ["Path", "Source", "Used For", "API Prefix"],
        [
            ["Statistics", "Postgres SQL aggregates", "Min/max/avg prices, charts, bedroom breakdown", "/stats"],
            ["ML Predictions", "joblib model file", "Price estimates, risk scores", "/predictions"],
            ["Property Data", "Postgres rows", "Search, listings, favorites, agent CRUD", "/properties"],
        ],
    )

    # 13. Environment Variables
    add_heading(doc, "13. Environment Variables", 1)
    add_table(
        doc,
        ["Variable", "Location", "Description"],
        [
            ["DATABASE_URL", "backend/.env", "Neon pooled PostgreSQL connection string (SSL required)"],
            ["SECRET_KEY", "backend/.env", "JWT signing secret"],
            ["CORS_ORIGINS", "backend/.env", "Allowed frontend origins (comma-separated)"],
            ["SEED_LIMIT", "backend/.env", "Optional: limit rows seeded (0 = all)"],
            ["VITE_API_URL", "frontend/.env", "Backend API URL (default: proxied /api)"],
        ],
    )

    # 14. Troubleshooting
    add_heading(doc, "14. Common Issues & Solutions", 1)
    add_table(
        doc,
        ["Issue", "Solution"],
        [
            ["alembic/uvicorn not recognized", "Use python -m alembic and python -m uvicorn on Windows"],
            [".env not found", "Place .env in backend/ folder or project root"],
            ["Password authentication failed", "Reset password in Neon dashboard, copy fresh pooled URL"],
            ["relation users does not exist", "Run python -m alembic upgrade head before seed_db"],
            ["404 on http://localhost:8000/", "Normal — use /docs or run frontend on :5173"],
            ["Model metrics not found", "Run python -m app.ml.train"],
        ],
    )

    # Footer
    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_para(doc, f"Generated: {date.today().strftime('%B %d, %Y')}")
    add_para(doc, "Project: PropAI.pk — Predictive AI for Real Estate Investment Analysis")
    add_para(doc, "Scope: Karachi & Islamabad, Pakistan")
    add_para(doc, "Repository structure: /backend (FastAPI) + /frontend (React/Vite)")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
